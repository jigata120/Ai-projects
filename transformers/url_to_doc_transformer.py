from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException, WebDriverException
from bs4 import BeautifulSoup
from docx import Document
from urllib.parse import urlparse
from webdriver_manager.chrome import ChromeDriverManager


chrome_options = Options()
chrome_options.add_argument("--headless")  
chrome_options.add_argument("--no-sandbox")
chrome_options.add_argument("--disable-dev-shm-usage")
service = Service(ChromeDriverManager().install())   
driver = webdriver.Chrome(service=service, options=chrome_options)

def extract_main_url_name(url):
    parsed_url = urlparse(url)
    domain = parsed_url.netloc
    domain_parts = domain.split('.')
    if len(domain_parts) >= 2:
        return '.'.join(domain_parts[-2:])
    else:
        return domain

def transform_to_doc(url):
    try:
        driver.get(url)
        WebDriverWait(driver, 10).until(lambda d: d.execute_script('return document.readyState') == 'complete')

        WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.CSS_SELECTOR, "div.content")))

        page_source = driver.page_source

    except TimeoutException:
        print("Timeout: Element or page not loaded in time")
        page_source = driver.page_source 
    except WebDriverException as e:
        print(f"WebDriverException occurred: {e}")
        return
    finally:
        driver.quit()

    with open('page_source.html', 'w', encoding='utf-8') as f:
        f.write(page_source)

    soup = BeautifulSoup(page_source, 'html.parser')

    doc = Document()

    def extract_and_add_text(element, doc):
        if element.name == 'h1':
            doc.add_heading(element.get_text(), level=1)
        elif element.name == 'h2':
            doc.add_heading(element.get_text(), level=2)
        elif element.name == 'h3':
            doc.add_heading(element.get_text(), level=3)
        elif element.name == 'ul' or element.name == 'ol':
            for li in element.find_all('li'):
                doc.add_paragraph(li.get_text(), style='ListBullet')
        else:
            if element.get_text(strip=True):
                doc.add_paragraph(element.get_text(strip=True))

    for element in soup.find_all(['h1', 'h2', 'h3', 'ul', 'ol', 'p']):
        extract_and_add_text(element, doc)

    filename = f'{extract_main_url_name(url)}.docx'
    doc.save(filename)
transform_to_doc('https://www.speedtest.net/result/16676933624')